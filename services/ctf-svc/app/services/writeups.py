"""Event writeups: upload, turn-in, and who is eliminated for not sending one."""

from __future__ import annotations

import io
import uuid
from datetime import datetime, timezone
from typing import Any
from uuid import UUID

from minio import Minio
from minio.error import S3Error
from sqlalchemy import and_, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import Settings, get_settings
from app.core.errors import AppError, ErrorCode
from app.core.logging import get_logger
from app.models import Event, EventWriteup

log = get_logger("writeups")

#: Extension → content type. Deliberately a small, closed set.
#:
#: Not a taste preference: an organiser has to read these in the browser without
#: downloading anything, and every format here can be shown. A .zip or a .html
#: could not be — and an uploaded .html rendered inline would run its own script
#: in the organiser's session.
ALLOWED: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".md": "text/markdown",
    ".txt": "text/plain",
}

MAX_BYTES = 25 * 1024 * 1024


class WriteupStorage:
    """The private bucket writeups live in.

    No public policy and no presigned URLs: every read goes through the service
    so the caller's role is checked each time. A presigned link would outlive
    the check and could be forwarded to anyone.
    """

    def __init__(self, settings: Settings) -> None:
        self._bucket = settings.storage_writeups_bucket
        self._region = settings.storage_region
        self._client = Minio(
            settings.storage_endpoint,
            access_key=settings.storage_access_key.get_secret_value(),
            secret_key=settings.storage_secret_key.get_secret_value(),
            secure=settings.storage_use_ssl,
            region=settings.storage_region,
        )

    def _ensure_bucket(self) -> None:
        if not self._client.bucket_exists(self._bucket):
            self._client.make_bucket(self._bucket, location=self._region)
            log.info("writeup_bucket_created", bucket=self._bucket)

    def put(self, key: str, data: bytes, content_type: str) -> None:
        self._ensure_bucket()
        self._client.put_object(
            self._bucket, key, io.BytesIO(data), length=len(data), content_type=content_type
        )

    def get(self, key: str) -> bytes:
        try:
            response = self._client.get_object(self._bucket, key)
            try:
                return response.read()
            finally:
                response.close()
                response.release_conn()
        except S3Error as exc:
            raise AppError(ErrorCode.NOT_FOUND, "the stored file is missing") from exc

    def delete(self, key: str) -> None:
        try:
            self._client.remove_object(self._bucket, key)
        except S3Error:
            # The row is the record of truth; a missing object should not stop
            # the captain replacing a broken upload.
            log.warning("writeup_object_missing_on_delete", key=key)


_storage: WriteupStorage | None = None


def get_storage() -> WriteupStorage:
    global _storage
    if _storage is None:
        _storage = WriteupStorage(get_settings())
    return _storage


class WriteupService:
    def __init__(self, session: AsyncSession) -> None:
        self.session = session

    # -- reads ---------------------------------------------------------------

    async def for_entry(
        self, event_id: UUID, *, team_id: UUID | None, user_id: UUID | None
    ) -> EventWriteup | None:
        where = (
            EventWriteup.team_id == team_id
            if team_id is not None
            else EventWriteup.user_id == user_id
        )
        row = await self.session.execute(
            select(EventWriteup).where(and_(EventWriteup.event_id == event_id, where))
        )
        return row.scalar_one_or_none()

    async def get(self, writeup_id: UUID) -> EventWriteup:
        row = await self.session.execute(
            select(EventWriteup).where(EventWriteup.id == writeup_id)
        )
        writeup = row.scalar_one_or_none()
        if not writeup:
            raise AppError(ErrorCode.NOT_FOUND, "writeup not found")
        return writeup

    async def list_for_event(self, event_id: UUID) -> list[EventWriteup]:
        rows = await self.session.execute(
            select(EventWriteup)
            .where(EventWriteup.event_id == event_id)
            .order_by(EventWriteup.updated_at.desc())
        )
        return list(rows.scalars())

    # -- writes --------------------------------------------------------------

    async def upload(
        self,
        event: Event,
        *,
        team_id: UUID | None,
        user_id: UUID | None,
        uploader_id: UUID,
        filename: str,
        data: bytes,
    ) -> EventWriteup:
        """Store a writeup, replacing any draft this entry already has.

        Refused once the entry has turned theirs in: submitting is the act that
        makes it count, and a file that can still change afterwards would make
        the deadline meaningless.
        """
        if not data:
            raise AppError(ErrorCode.VALIDATION, "the file is empty")
        if len(data) > MAX_BYTES:
            raise AppError(
                ErrorCode.VALIDATION, f"the file is larger than {MAX_BYTES // (1024 * 1024)} MB"
            )

        suffix = _extension(filename)
        content_type = ALLOWED.get(suffix)
        if content_type is None:
            raise AppError(
                ErrorCode.VALIDATION,
                f"unsupported file type {suffix or '(none)'} — allowed: "
                + ", ".join(sorted(ALLOWED)),
            )

        existing = await self.for_entry(event.id, team_id=team_id, user_id=user_id)
        if existing is not None and existing.status == "submitted":
            raise AppError(
                ErrorCode.VALIDATION,
                "this writeup has been turned in — it cannot be replaced",
            )

        storage = get_storage()
        key = f"{event.id}/{team_id or user_id}/{uuid.uuid4().hex}{suffix}"
        storage.put(key, data, content_type)

        if existing is not None:
            # Replacing a draft: drop the old object so the bucket does not
            # accumulate files nothing points at.
            old_key = existing.storage_key
            existing.filename = filename
            existing.content_type = content_type
            existing.size_bytes = len(data)
            existing.storage_key = key
            existing.uploaded_by = uploader_id
            existing.updated_at = datetime.now(timezone.utc)
            await self.session.flush()
            storage.delete(old_key)
            log.info("writeup_replaced", event_id=str(event.id), writeup_id=str(existing.id))
            return existing

        writeup = EventWriteup(
            event_id=event.id,
            team_id=team_id,
            user_id=user_id,
            filename=filename,
            content_type=content_type,
            size_bytes=len(data),
            storage_key=key,
            status="draft",
            uploaded_by=uploader_id,
        )
        self.session.add(writeup)
        await self.session.flush()
        log.info("writeup_uploaded", event_id=str(event.id), writeup_id=str(writeup.id))
        return writeup

    async def remove(self, writeup: EventWriteup) -> None:
        """Delete a draft. A turned-in writeup stays: it is a submission."""
        if writeup.status == "submitted":
            raise AppError(
                ErrorCode.VALIDATION, "a writeup that has been turned in cannot be removed"
            )
        key = writeup.storage_key
        await self.session.delete(writeup)
        await self.session.flush()
        get_storage().delete(key)
        log.info("writeup_removed", writeup_id=str(writeup.id))

    async def turn_in(self, event: Event, writeup: EventWriteup) -> EventWriteup:
        """Mark it submitted. This is the act the deadline measures."""
        if writeup.status == "submitted":
            raise AppError(ErrorCode.VALIDATION, "already turned in")

        now = datetime.now(timezone.utc)
        if event.writeup_deadline and now > event.writeup_deadline:
            raise AppError(ErrorCode.VALIDATION, "the writeup deadline has passed")

        writeup.status = "submitted"
        writeup.submitted_at = now
        writeup.updated_at = now
        await self.session.flush()
        log.info("writeup_turned_in", writeup_id=str(writeup.id))
        return writeup

    # -- preview -------------------------------------------------------------

    def render(self, writeup: EventWriteup) -> tuple[bytes, str, str]:
        """Bytes to show in the browser, its content type, and how to show it.

        The third value tells the client what it is dealing with: `pdf` goes in
        an embed, `text` and `markdown` are rendered as themselves, and a .docx
        is unzipped here into plain text because nothing in a browser reads one.
        """
        raw = get_storage().get(writeup.storage_key)

        if writeup.content_type == "application/pdf":
            return raw, writeup.content_type, "pdf"

        if writeup.filename.lower().endswith(".docx"):
            return _docx_to_text(raw).encode("utf-8"), "text/plain; charset=utf-8", "text"

        kind = "markdown" if writeup.filename.lower().endswith(".md") else "text"
        return raw, "text/plain; charset=utf-8", kind


def _extension(filename: str) -> str:
    name = (filename or "").lower().strip()
    dot = name.rfind(".")
    return name[dot:] if dot != -1 else ""


def _docx_to_text(raw: bytes) -> str:
    """Paragraphs and table cells, in order. Enough to read and judge by."""
    try:
        import docx  # imported lazily: only .docx uploads pay for it
    except ImportError:  # pragma: no cover
        return "(cannot preview .docx on this server)"

    document = docx.Document(io.BytesIO(raw))
    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip() or "(the document has no readable text)"


def required_ranks(event: Event) -> int | None:
    """How far down the board the writeup requirement reaches."""
    return event.writeup_required_top_n


def eliminated_now(
    event: Event,
    *,
    ranked: list[dict[str, Any]],
    submitted_subjects: set[str],
) -> set[str]:
    """Which entries are out for not turning in a writeup.

    Computed, never stored: it becomes true when the clock passes the deadline
    and stops being true the moment an organiser extends it or a writeup lands.
    A stored flag would have to be kept in step with both and would be wrong in
    between.

    `ranked` is the board *before* elimination, which is what decides who owed a
    writeup. Deciding it afterwards would be circular — eliminating third place
    promotes fourth into the top three, who would then owe one they were never
    asked for.
    """
    top_n = event.writeup_required_top_n
    if not top_n or not event.writeup_deadline:
        return set()
    if datetime.now(timezone.utc) <= event.writeup_deadline:
        return set()

    owed = {row["subject"] for row in ranked[:top_n]}
    return owed - submitted_subjects
